from __future__ import annotations

import io
import os
import tempfile
from pathlib import Path

from flask import Flask, render_template, request, send_file
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from werkzeug.utils import secure_filename
from docx import Document

app = Flask(__name__)
ALLOWED_EXTENSIONS = {"txt", "md", "docx", "pdf"}
TARGET_FORMATS = {"pdf", "word", "markdown"}


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".docx":
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    if suffix == ".pdf":
        reader = PdfReader(str(file_path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    raise ValueError("不支持的源文件格式")


def to_markdown(text: str) -> io.BytesIO:
    output = io.BytesIO()
    output.write(text.encode("utf-8"))
    output.seek(0)
    return output


def to_word(text: str) -> io.BytesIO:
    doc = Document()
    for line in text.splitlines() or [text]:
        doc.add_paragraph(line)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def to_pdf(text: str) -> io.BytesIO:
    output = io.BytesIO()
    c = canvas.Canvas(output, pagesize=A4)
    width, height = A4
    y = height - 40

    for line in text.splitlines() or [text]:
        wrapped_lines = [line[i : i + 70] for i in range(0, len(line), 70)] or [""]
        for wrapped in wrapped_lines:
            c.drawString(40, y, wrapped)
            y -= 18
            if y < 40:
                c.showPage()
                y = height - 40

    c.save()
    output.seek(0)
    return output


@app.get("/")
def index():
    return render_template("index.html")


@app.post("/convert")
def convert_file():
    uploaded = request.files.get("file")
    target = request.form.get("target", "").lower()

    if uploaded is None or uploaded.filename == "":
        return "请选择要上传的文件。", 400

    if not allowed_file(uploaded.filename):
        return "文件格式不支持，仅支持 txt/md/docx/pdf。", 400

    if target not in TARGET_FORMATS:
        return "目标格式不支持，仅支持 pdf/word/markdown。", 400

    filename = secure_filename(uploaded.filename)

    with tempfile.TemporaryDirectory() as tmp_dir:
        source_path = Path(tmp_dir) / filename
        uploaded.save(source_path)

        text = extract_text(source_path)

    stem = Path(filename).stem
    if target == "pdf":
        data = to_pdf(text)
        out_name = f"{stem}.pdf"
        mime = "application/pdf"
    elif target == "word":
        data = to_word(text)
        out_name = f"{stem}.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        data = to_markdown(text)
        out_name = f"{stem}.md"
        mime = "text/markdown"

    return send_file(data, as_attachment=True, download_name=out_name, mimetype=mime)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)
